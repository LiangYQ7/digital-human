import os
from pathlib import Path

# Verify python-docx is available
try:
    from docx import Document
    print("python-docx OK")
except ImportError:
    print("python-docx NOT AVAILABLE")
    raise

base = Path(r"D:\20260323113204906\示范景区公开资料包")
outdir = Path(r"D:\Code\Vs code\digital_human\brain\data\raw")
outdir.mkdir(parents=True, exist_ok=True)

docx_files = [
    "灵山胜境 景点结构化数据集.docx",
    "灵山胜境：历史、文化、景点特色与个性化游览指南.docx",
]

for fname in docx_files:
    path = base / fname
    if not path.exists():
        print(f"MISSING: {path}")
        continue
    print(f"Processing: {fname}")
    doc = Document(str(path))
    
    # Extract all text (paragraphs + tables)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {ti} ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    
    full_text = "\n".join(lines)
    
    # Save as .txt in brain/data/raw
    txt_name = fname.replace(".docx", ".txt")
    txt_path = outdir / txt_name
    txt_path.write_text(full_text, encoding="utf-8")
    print(f"Written {len(full_text)} chars to {txt_path}")
    
    # Print first 2000 chars for verification
    print("--- PREVIEW ---")
    print(full_text[:2000])
    print("--- END PREVIEW ---")

print("DONE")
