"""Pytest test that extracts full text from two docx files."""
import sys
import os

OUTPUT_PATH = r"D:\Code\Vs code\digital_human\_docx_full_extract.txt"

def write_report(text):
    with open(OUTPUT_PATH, 'a', encoding='utf-8') as f:
        f.write(text + '\n')

def test_extract_docx_files():
    """Extract all text from two docx files and write to output."""
    import docx as docx_lib
    
    docx_files = [
        (r'D:\20260323113204906\示范景区公开资料包\灵山胜境 景点结构化数据集.docx', 'FILE_1'),
        (r'D:\20260323113204906\示范景区公开资料包\灵山胜境：历史、文化、景点特色与个性化游览指南.docx', 'FILE_2'),
    ]
    
    # Clear previous output
    if os.path.exists(OUTPUT_PATH):
        os.remove(OUTPUT_PATH)
    
    for filepath, label in docx_files:
        fname = os.path.basename(filepath)
        write_report("=" * 80)
        write_report(f"LABEL: {label}")
        write_report(f"FILENAME: {fname}")
        write_report(f"FULL PATH: {filepath}")
        write_report("=" * 80)
        
        doc = docx_lib.Document(filepath)
        
        # Collect all paragraphs
        all_paragraphs = []
        non_empty_count = 0
        headings = []
        total_chars_nonempty = 0
        
        for para in doc.paragraphs:
            text = para.text
            all_paragraphs.append(text)
            if text.strip():
                non_empty_count += 1
                total_chars_nonempty += len(text)
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name or "heading" in style_name:
                headings.append((style_name, text[:150]))
        
        table_count = len(doc.tables)
        full_text = "\n".join(all_paragraphs)
        total_chars_full = len(full_text)
        
        # METADATA
        write_report("")
        write_report("--- METADATA ---")
        write_report(f"Total paragraphs (including empty): {len(all_paragraphs)}")
        write_report(f"Non-empty paragraphs: {non_empty_count}")
        write_report(f"Characters in non-empty paragraphs: {total_chars_nonempty}")
        write_report(f"Total characters (full text, including whitespace): {total_chars_full}")
        write_report(f"Tables: {table_count}")
        write_report(f"Headings: {len(headings)}")
        
        # STRUCTURE
        if headings:
            write_report("")
            write_report("--- STRUCTURE (HEADINGS) ---")
            for style, text in headings:
                write_report(f"  [{style}] {text}")
        
        # FULL TEXT
        write_report("")
        write_report("--- FULL TEXT (START) ---")
        write_report(full_text)
        write_report("--- FULL TEXT (END) ---")
        write_report("")
        
        # TABLE CONTENT
        if table_count > 0:
            write_report("")
            write_report("--- TABLE CONTENTS ---")
            for ti, table in enumerate(doc.tables):
                write_report(f"Table {ti+1}: {len(table.rows)} rows x {len(table.columns)} cols")
                for ri, row in enumerate(table.rows):
                    cells_text = [cell.text for cell in row.cells]
                    write_report(f"  Row {ri}: {cells_text}")
                write_report("")
    
    write_report("=" * 80)
    write_report("EXTRACTION COMPLETE")
    write_report("=" * 80)
    
    assert True
