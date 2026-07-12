import sys, os, traceback

output_path = r"D:\Code\Vs code\digital_human\_docx_extract_result.txt"

try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"],
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    import docx

results = []

files = [
    (r"D:\20260323113204906\示范景区公开资料包\灵山胜境 景点结构化数据集.docx", "File1"),
    (r"D:\20260323113204906\示范景区公开资料包\灵山胜境：历史、文化、景点特色与个性化游览指南.docx", "File2"),
]

for filepath, label in files:
    if not os.path.exists(filepath):
        results.append(f"=== {label}: FILE NOT FOUND at {filepath} ===\n")
        continue
        
    doc = docx.Document(filepath)
    
    all_paragraphs = []
    non_empty_count = 0
    headings = []
    total_chars = 0
    
    for para in doc.paragraphs:
        text = para.text
        all_paragraphs.append(text)
        if text.strip():
            non_empty_count += 1
            total_chars += len(text)
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name or "heading" in style_name:
            headings.append((style_name, text[:120]))
    
    table_count = len(doc.tables)
    full_text = "\n".join(all_paragraphs)
    
    result = f"""
{'='*80}
FILE: {label}
PATH: {filepath}
{'='*80}

--- METADATA ---
Non-empty paragraphs: {non_empty_count}
Total characters (non-empty paragraphs): {total_chars}
Total characters (full text incl whitespace): {len(full_text)}
Tables found: {table_count}
Headings detected: {len(headings)}

--- STRUCTURE (Headings) ---
"""
    for style, text in headings:
        result += f"  [{style}] {text}\n"
    
    result += f"""
--- FULL TEXT ---
{full_text}

--- END OF FILE ---
"""
    results.append(result)

final_output = "\n".join(results)

with open(output_path, 'w', encoding='utf-8') as f:
    f.write(final_output)

# Also try to print to stdout
print("EXTRACTION COMPLETE")
print(f"Output written to: {output_path}")
print(f"Total result size: {len(final_output)} chars")
