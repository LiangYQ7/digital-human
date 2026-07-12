"""读取赛题资料包中的 docx 文件，输出文本内容"""
import os, sys
from pathlib import Path

base = Path(r"D:\20260323113204906\示范景区公开资料包")

for fname in [
    "灵山胜境 景点结构化数据集.docx",
    "灵山胜境：历史、文化、景点特色与个性化游览指南.docx",
]:
    path = base / fname
    print(f"\n===== {fname} =====", flush=True)
    try:
        from docx import Document
        doc = Document(str(path))
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                print(f"[P{i}] {text[:300]}", flush=True)
        
        # Also check tables
        for ti, table in enumerate(doc.tables):
            print(f"[TABLE{ti}] rows={len(table.rows)}, cols={len(table.columns)}", flush=True)
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip()[:80] for cell in row.cells]
                print(f"  R{ri}: {' | '.join(cells)}", flush=True)
                if ri > 10:
                    print(f"  ... ({len(table.rows) - 11} more rows)", flush=True)
                    break
    except ImportError:
        print("python-docx not installed, trying zipfile...", flush=True)
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(str(path)) as z:
            xml = z.read("word/document.xml")
            root = ET.fromstring(xml)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for p in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
                texts = [t.text for t in p.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t") if t.text]
                if texts:
                    line = "".join(texts)
                    if line.strip():
                        print(line.strip()[:300], flush=True)
