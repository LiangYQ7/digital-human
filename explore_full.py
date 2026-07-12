"""
Full exploration of the scenic area data package.
Run: "D:/conda_data/envs/fay/python" explore_full.py
Output: _full_report.txt
"""
import sys, os

OUT = r"D:\Code\Vs code\digital_human\_full_report.txt"

def log(s):
    print(s, flush=True)
    with open(OUT, 'a', encoding='utf-8') as f:
        f.write(s + '\n')

# Clear previous report
if os.path.exists(OUT):
    os.remove(OUT)

log("=" * 70)
log("FULL EXPLORATION REPORT — 示范景区公开资料包")
log("=" * 70)

# ============================================================================
# PART 1: Excel file
# ============================================================================
log("\n" + "=" * 70)
log("PART 1: EXCEL FILE — 景点景区旅游数据行为分析数据.xlsx")
log("=" * 70)

import openpyxl
import pandas as pd

xlsx_path = r'D:\20260323113204906\示范景区公开资料包\景点景区旅游数据行为分析数据.xlsx'

# 1a. Sheet names
log("\n--- 1a. Sheet Names ---")
wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
sheets = list(wb.sheetnames)
for i, name in enumerate(sheets):
    log(f"  Sheet [{i}]: '{name}'")
wb.close()
log(f"  Total: {len(sheets)} sheet(s)")

# 1b. Per-sheet: columns, row count, first 5 rows
for i, name in enumerate(sheets):
    log(f"\n--- 1b. Sheet '{name}' ---")
    try:
        df = pd.read_excel(xlsx_path, sheet_name=name, nrows=None)
        total_rows = len(df)
        cols = list(df.columns)
        log(f"  Total rows: {total_rows}")
        log(f"  Columns ({len(cols)}):")
        for c in cols:
            log(f"    - {c}")
        log(f"\n  First 5 rows:")
        for line in df.head(5).to_string().split('\n'):
            log(f"    {line}")
        log(f"\n  Last 3 rows:")
        for line in df.tail(3).to_string().split('\n'):
            log(f"    {line}")
        # Data types
        log(f"\n  Dtypes:")
        for col, dtype in df.dtypes.items():
            log(f"    {col}: {dtype}")
        # Null counts
        nulls = df.isnull().sum()
        if nulls.sum() > 0:
            log(f"\n  Null counts:")
            for col, cnt in nulls.items():
                if cnt > 0:
                    log(f"    {col}: {cnt} nulls")
    except Exception as e:
        log(f"  ERROR: {e}")

# 1c. Keyword search: 灵山胜境, 拈花湾
log(f"\n--- 1c. Keyword Search ---")
keywords = ['灵山胜境', '拈花湾']

for name in sheets:
    try:
        df = pd.read_excel(xlsx_path, sheet_name=name, dtype=str)
        log(f"\n  Sheet '{name}':")
        for kw in keywords:
            mask = df.apply(lambda row: row.astype(str).str.contains(kw, na=False).any(), axis=1)
            count = mask.sum()
            log(f"    '{kw}': {count} matching rows out of {len(df)}")
            if count > 0:
                sample = df[mask]
                log(f"    Sample (first 3):")
                for line in sample.head(3).to_string().split('\n'):
                    log(f"      {line}")
                # Show which columns matched
                log(f"    Columns where '{kw}' appears:")
                for col in df.columns:
                    col_matches = df[col].astype(str).str.contains(kw, na=False).sum()
                    if col_matches > 0:
                        log(f"      {col}: {col_matches} matches")
    except Exception as e:
        log(f"  ERROR in sheet '{name}': {e}")

# 1d. Filterable columns summary
log(f"\n--- 1d. Filterable Columns Summary ---")
for name in sheets:
    try:
        df = pd.read_excel(xlsx_path, sheet_name=name, nrows=0)
        log(f"\n  Sheet '{name}' columns:")
        for col in df.columns:
            log(f"    - {col}")
    except Exception as e:
        log(f"  ERROR: {e}")

# ============================================================================
# PART 2: Docx files
# ============================================================================
log("\n" + "=" * 70)
log("PART 2: DOCX FILES")
log("=" * 70)

import docx as docx_lib

docx_files = [
    r'D:\20260323113204906\示范景区公开资料包\灵山胜境 景点结构化数据集.docx',
    r'D:\20260323113204906\示范景区公开资料包\灵山胜境：历史、文化、景点特色与个性化游览指南.docx',
]

for docx_path in docx_files:
    fname = os.path.basename(docx_path)
    log(f"\n--- File: {fname} ---")
    try:
        doc = docx_lib.Document(docx_path)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        # Count total characters
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)
        
        log(f"  Paragraphs: {para_count}")
        log(f"  Tables: {table_count}")
        log(f"  Total characters: {total_chars}")
        
        # Show structure (first 500 chars of each paragraph that has text)
        log(f"\n  Content preview:")
        shown = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                prefix = text[:200]
                if len(text) > 200:
                    prefix += "..."
                log(f"    [{p.style.name}] {prefix}")
                shown += 1
                if shown >= 50:
                    log(f"    ... (truncated, {para_count - shown} more paragraphs)")
                    break
        
        # Show table previews
        if table_count > 0:
            log(f"\n  Table previews:")
            for ti, table in enumerate(doc.tables):
                log(f"    Table {ti+1}: {len(table.rows)} rows x {len(table.columns)} cols")
                # Show header + first 3 rows
                for ri, row in enumerate(table.rows):
                    if ri > 3:
                        log(f"      ... ({len(table.rows) - 4} more rows)")
                        break
                    cells_text = [cell.text.strip()[:50] for cell in row.cells]
                    log(f"      Row {ri}: {cells_text}")
    except Exception as e:
        log(f"  ERROR: {e}")

log("\n" + "=" * 70)
log("REPORT COMPLETE")
log("=" * 70)
print("\nDONE! Report saved to:", OUT)
